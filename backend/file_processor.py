# File Processor Module
# Handles PDF and Word document uploads for Resume Evaluator

import os
from typing import Tuple, Optional
import traceback


async def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text content from a PDF file."""
    try:
        import io
        from PyPDF2 import PdfReader
        
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return "\n".join(text_parts)
    except ImportError:
        raise ImportError("PyPDF2 is required for PDF processing. Install with: pip install PyPDF2")
    except Exception as e:
        print(f"[ERROR] Failed to extract text from PDF: {str(e)}")
        traceback.print_exc()
        raise ValueError(f"Failed to process PDF file: {str(e)}")


async def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text content from a Word document (.docx)."""
    try:
        import io
        from docx import Document
        
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    except ImportError:
        raise ImportError("python-docx is required for Word document processing. Install with: pip install python-docx")
    except Exception as e:
        print(f"[ERROR] Failed to extract text from DOCX: {str(e)}")
        traceback.print_exc()
        raise ValueError(f"Failed to process Word document: {str(e)}")


async def process_resume_file(filename: str, file_content: bytes) -> Tuple[str, str]:
    """
    Process uploaded resume file and extract text.
    
    Args:
        filename: Original filename with extension
        file_content: Binary content of the file
        
    Returns:
        Tuple of (extracted_text, file_type)
    """
    if not filename or not file_content:
        raise ValueError("No file provided")
    
    # Get file extension
    _, ext = os.path.splitext(filename.lower())
    
    if ext == ".pdf":
        text = await extract_text_from_pdf(file_content)
        return text, "pdf"
    elif ext in [".docx", ".doc"]:
        if ext == ".doc":
            raise ValueError("Legacy .doc format is not supported. Please convert to .docx or PDF.")
        text = await extract_text_from_docx(file_content)
        return text, "docx"
    else:
        raise ValueError(f"Unsupported file format: {ext}. Please upload a PDF or Word document (.docx)")


def validate_file_size(file_content: bytes, max_size_mb: int = 5) -> bool:
    """Check if file size is within allowed limit."""
    max_bytes = max_size_mb * 1024 * 1024
    return len(file_content) <= max_bytes


def get_allowed_extensions() -> list:
    """Get list of allowed file extensions for resume upload."""
    return [".pdf", ".docx"]
